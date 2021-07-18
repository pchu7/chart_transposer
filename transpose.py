# Chart Transposer
# Author: Phillip Chu
# Date: 7/7/2021
# Description: A script to transpose the key of a chord chart 
# stored as a word document

# ----------------------
# General Includes
# ----------------------
import docx


# ----------------------
# Variables 
# ----------------------

# ----------------------
# Procedures
# ----------------------

# ----------------------
# proc_input 
# ----------------------
#def proc_input():
# ----------------------
# __main__
# ----------------------

def main():
    doc = docx.Document("Living Hope - Test0.docx")
    transposed_doc = docx.Document()
    all_paragraphs = doc.paragraphs
    print ( len ( all_paragraphs ) )
    for para in all_paragraphs:
        #print ( para.text )
        runs = para.runs
        for run in runs:
            if ( run.bold == True and run.italic == True ):
                print ( run.text ) 
            

if __name__ == "__main__":
    main()
